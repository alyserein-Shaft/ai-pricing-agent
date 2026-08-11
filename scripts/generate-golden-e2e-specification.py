#!/usr/bin/env python3
"""Generate the deterministic, native-text Golden E2E technical reference."""

from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUTPUT = ROOT / "tests/e2e/fixtures/golden-specification.docx"


def main() -> None:
    document = Document()
    section = document.sections[0]
    section.top_margin = section.bottom_margin = Inches(1)
    section.left_margin = section.right_margin = Inches(1)

    normal = document.styles["Normal"]
    normal.font.name = "Arial"
    normal.font.size = Pt(10.5)
    normal.paragraph_format.space_after = Pt(6)

    for style_name, size in (("Title", 22), ("Heading 1", 16), ("Heading 2", 12)):
        style = document.styles[style_name]
        style.font.name = "Arial"
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor(31, 78, 121)

    document.add_heading("Golden E2E Technical Reference", 0)
    document.add_paragraph(
        "Deterministic native-text fixture for the governed AI Pricing Agent journey. "
        "No OCR is required."
    )
    document.add_heading("1. Scope", level=1)
    document.add_paragraph("Fire alarm equipment for the Golden E2E validation project.")
    document.add_heading("2. Technical requirements", level=1)
    document.add_heading("2.1 Addressable detector", level=2)
    document.add_paragraph(
        "Provide an addressable detector by Golden Manufacturer, exact model "
        "GOLDEN-FA-001, operating at 24 V DC; it shall comply with UL 268, "
        "shall be compatible with Golden Fire Addressable Control Panel GF-CP-001, "
        "and shall include a detector base."
    )
    document.add_heading("2.2 Addressable interface module", level=2)
    document.add_paragraph(
        "Provide an addressable interface module. The manufacturer and exact model "
        "are not stated and require clarification."
    )
    document.add_heading("2.3 Specialized annunciator", level=2)
    document.add_paragraph(
        "Provide a specialized unsupported annunciator by Golden Manufacturer, exact "
        "model GOLDEN-NOMATCH-001."
    )
    document.add_heading("3. Evidence control", level=1)
    document.add_paragraph(
        "Only explicit evidence stated above is authoritative. Compatibility, compliance, "
        "and detector-base evidence for GOLDEN-FA-001 are intentional positive-path evidence. "
        "Missing identity, compatibility, certification, accessory, and commercial evidence "
        "for the incomplete interface module and unsupported annunciator must not be inferred."
    )

    document.core_properties.title = "Golden E2E Technical Reference"
    document.core_properties.subject = "Deterministic test fixture"
    document.core_properties.author = "AI Pricing Agent Golden E2E"
    document.save(OUTPUT)


if __name__ == "__main__":
    main()
